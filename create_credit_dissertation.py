from docx import Document

doc = Document()

doc.add_heading('Credit Score System Dissertation', level=1)

doc.add_heading('Chapter 7: Design', level=1)

doc.add_heading('7.1 Introduction', level=2)
doc.add_paragraph(
    'This chapter outlines the design of the proposed Credit Score System. ' 
    'It explains how the system will function, describes the architecture and ' 
    'physical deployment, presents the database and program design, and ' 
    'defines the user interface, security controls, and supporting algorithms.')

# System Design

doc.add_heading('7.2 System Design', level=2)
doc.add_paragraph(
    'The Credit Score System is designed to collect borrower financial data, ' 
    'evaluate credit risk using predefined scoring rules, and generate a reliable ' 
    'credit score for use by lenders and financial institutions. The system ' 
    'supports borrower registration, data entry, risk assessment, score review, ' 
    'and report generation.')

doc.add_paragraph('Key system functions include:')
doc.add_paragraph('• Borrower profile management', style='List Bullet')
doc.add_paragraph('• Financial data collection and validation', style='List Bullet')
doc.add_paragraph('• Credit score calculation and scoring logic', style='List Bullet')
doc.add_paragraph('• Report generation and decision support', style='List Bullet')

doc.add_heading('7.2.1 Context Diagram and DFD of the proposed System', level=3)
doc.add_paragraph(
    'The context diagram shows the Credit Score System as a single process with ' 
    'external entities including Borrower, Credit Analyst, and Third-Party Data Sources. ' 
    'The system exchanges borrower information, credit scores, and reports with these entities.')
doc.add_paragraph(
    'The data flow diagram (DFD) levels include:')
doc.add_paragraph('• Level 0: System input of borrower credentials, financial records, and credit bureau data.', style='List Bullet')
doc.add_paragraph('• Level 1: Processes for data validation, score calculation, and report generation.', style='List Bullet')
doc.add_paragraph('• Level 2: Subprocesses for risk factor computation, data storage, and feedback.', style='List Bullet')
doc.add_paragraph(
    'In the completed dissertation, include diagrams created using a diagram tool and ' 
    'insert them in this section to show system boundaries and data flows clearly.')

# Architectural Design

doc.add_heading('7.3 Architectural design', level=2)
doc.add_paragraph(
    'The system architecture follows a three-tier structure: presentation, application, ' 
    'and data storage. The presentation layer is a web-based user interface accessible ' 
    'to borrowers and analysts. The application layer handles business logic, including ' 
    'score calculations, validation, and report generation. The data layer provides ' 
    'persistent storage for borrower profiles, financial records, and audit logs.')
doc.add_paragraph('Architecture components:')
doc.add_paragraph('• User Interface Layer', style='List Bullet')
doc.add_paragraph('• Business Logic Layer', style='List Bullet')
doc.add_paragraph('• Database Layer', style='List Bullet')
doc.add_paragraph('• External Data Integration Layer', style='List Bullet')

doc.add_paragraph(
    'The system uses RESTful web services to integrate with credit bureaus and external ' 
    'financial data providers, enabling secure data exchange and modular scaling.')

# Physical design

doc.add_heading('7.4 Physical design', level=2)
doc.add_paragraph(
    'The physical design describes how software and hardware interact. The system ' 
    'will run on a cloud-hosted application server with a web server front-end. The database ' 
    'resides on a managed relational database instance. End users access the system through ' 
    'web browsers on laptops, desktops, or mobile devices.')

doc.add_paragraph('Hardware and software interaction:')
doc.add_paragraph('• Client devices: browsers send HTTP requests to the web server.', style='List Bullet')
doc.add_paragraph('• Application server: processes requests, executes scoring algorithms, and returns responses.', style='List Bullet')
doc.add_paragraph('• Database server: stores borrower records, score histories, and audit trails.', style='List Bullet')
doc.add_paragraph('• Third-party services: external credit bureau APIs supply additional data through secure channels.', style='List Bullet')

doc.add_heading('7.5 Database design', level=2)
doc.add_paragraph(
    'The database stores borrower accounts, financial metrics, credit score results, ' 
    'and audit records. A relational schema is suitable for this system.')

doc.add_paragraph('Core tables include:')
doc.add_paragraph('• Borrower', style='List Bullet')
doc.add_paragraph('• FinancialRecord', style='List Bullet')
doc.add_paragraph('• CreditScore', style='List Bullet')
doc.add_paragraph('• ScoreRule', style='List Bullet')
doc.add_paragraph('• AuditLog', style='List Bullet')

doc.add_paragraph(
    'An EER diagram should show entities and relationships such as Borrower 1-to-many ' 
    'FinancialRecord, Borrower 1-to-many CreditScore, and ScoreRule used in score calculations.')

doc.add_heading('7.6 Program design', level=2)
doc.add_paragraph(
    'Program design organizes the code into packages and classes. The main modules include user management, ' 
    'data ingestion, score computation, reporting, and security.')
doc.add_paragraph('Example package structure:')
doc.add_paragraph('• auth', style='List Bullet')
doc.add_paragraph('• borrower', style='List Bullet')
doc.add_paragraph('• scoring', style='List Bullet')
doc.add_paragraph('• reporting', style='List Bullet')
doc.add_paragraph('• persistence', style='List Bullet')

doc.add_paragraph(
    'Sequence diagrams should describe borrower registration, score calculation, and report generation flows. ' 
    'Collaboration diagrams should show interactions between UI controllers, scoring services, and the database.')

doc.add_heading('7.7 Interface design', level=2)
doc.add_paragraph(
    'The interface design focuses on usability and clear navigation for borrowers and credit analysts. ' 
    'It includes menus, input forms, and output reports designed for straightforward operation.')

doc.add_heading('7.7.1 Menu design', level=3)
doc.add_paragraph(
    'The main application menu organizes system functions by user role and task category.')
doc.add_heading('7.7.1.1 Main menu', level=4)
doc.add_paragraph('Main menu options:')
doc.add_paragraph('• Dashboard', style='List Bullet')
doc.add_paragraph('• Borrower Management', style='List Bullet')
doc.add_paragraph('• Credit Scoring', style='List Bullet')
doc.add_paragraph('• Reports', style='List Bullet')
doc.add_paragraph('• Administration', style='List Bullet')
doc.add_heading('7.7.1.2 Sub-menus', level=4)
doc.add_paragraph('Sub-menus may include:')
doc.add_paragraph('• Add New Borrower', style='List Bullet')
doc.add_paragraph('• Upload Financial Data', style='List Bullet')
doc.add_paragraph('• View Score History', style='List Bullet')
doc.add_paragraph('• Generate Report', style='List Bullet')

doc.add_heading('7.7.2 Input design', level=3)
doc.add_paragraph(
    'Input design covers all forms used to capture borrower and financial information. ' 
    'The system validates all entries to reduce errors and ensure data integrity.')
doc.add_paragraph('Key input forms include:')
doc.add_paragraph('• Borrower registration form', style='List Bullet')
doc.add_paragraph('• Income and employment form', style='List Bullet')
doc.add_paragraph('• Loan history and liability form', style='List Bullet')
doc.add_paragraph('• Credit bureau data import form', style='List Bullet')

doc.add_heading('7.7.3 Output design', level=3)
doc.add_paragraph(
    'Output design includes reports and summaries produced by the system. Output should be clear, action-oriented, and suitable for decision makers.')
doc.add_paragraph('Key output reports include:')
doc.add_paragraph('• Credit score report', style='List Bullet')
doc.add_paragraph('• Borrower risk summary', style='List Bullet')
doc.add_paragraph('• Credit history analysis', style='List Bullet')
doc.add_paragraph('• Audit log report', style='List Bullet')

doc.add_heading('7.8 Pseudo Code', level=2)
doc.add_paragraph(
    'Pseudo code defines the main scoring algorithm and validation process.')
doc.add_paragraph(
    'Example:\n\n'
    '1. Receive borrower ID and financial data.\n'
    '2. Validate borrower profile and required fields.\n'
    '3. Retrieve historical credit and repayment data.\n'
    '4. Apply scoring rules to compute risk points.\n'
    '5. Aggregate points into a final credit score.\n'
    '6. Store score result and generate report.\n'
)

doc.add_heading('7.9 Security design', level=2)
doc.add_paragraph(
    'Security design ensures the Credit Score System maintains confidentiality, integrity, and availability of borrower data.')
doc.add_heading('7.9.1 Physical security', level=3)
doc.add_paragraph(
    'Physical security involves protecting server infrastructure and office workstations. Measures include controlled data center access, locked server rooms, and secure workstations with screen locks.')
doc.add_heading('7.9.2 Network security', level=3)
doc.add_paragraph(
    'Network security includes firewalls, secure VPN access, encryption of data in transit, and intrusion detection systems. The system uses HTTPS for all browser connections and secure API channels for external data.')
doc.add_heading('7.9.3 Operational security', level=3)
doc.add_paragraph(
    'Operational security covers user authentication, role-based access control, audit logging, and regular security reviews. User accounts are limited to approved staff, and sensitive actions require appropriate privileges.')

doc.add_heading('7.10 Conclusion', level=2)
doc.add_paragraph(
    'This design chapter establishes the structure and security of the proposed Credit Score System. ' 
    'It demonstrates how the system will collect and validate data, calculate credit scores, and provide secure access to users. ' 
    'The design balances usability, reliability, and protectiveness to support responsible credit decisions.')

# Chapter 8

doc.add_heading('Chapter 8: Implementation and Testing', level=1)
doc.add_heading('8.1 Introduction', level=2)
doc.add_paragraph(
    'This chapter presents the implementation approach, testing strategy, installation plan, maintenance recommendations, and future development possibilities for the Credit Score System.')

doc.add_heading('8.2 Coding', level=2)
doc.add_paragraph(
    'The system was coded using a modular design with separate components for user interface, business logic, and data persistence. ' 
    'The codebase implements secure authentication, data validation, score calculations, and report generation. ' 
    'Key programming concepts include object-oriented design, service abstraction, and reusable modules.')
doc.add_heading('8.3 Testing', level=2)
doc.add_paragraph(
    'Testing includes functional test cases for borrower registration, data validation, score calculation, and report generation. ' 
    'Security testing includes authentication checks, authorization tests, input validation, and protection against common vulnerabilities such as SQL injection and unauthorized access.')
doc.add_paragraph('Example test cases:')
doc.add_paragraph('• Register new borrower with valid data', style='List Bullet')
doc.add_paragraph('• Reject incomplete borrower registration form', style='List Bullet')
doc.add_paragraph('• Generate credit score for a borrower with full history', style='List Bullet')
doc.add_paragraph('• Deny access to restricted administration functions for standard users', style='List Bullet')
doc.add_paragraph(
    'Screenshots of test case results and security enforcement should be inserted into this section when available. These screenshots demonstrate how the system behaves under normal and restricted conditions.')

doc.add_heading('8.4 Installation', level=2)
doc.add_paragraph(
    'Installation covers deploying the application, configuring the database, and training users. The recommended approach includes user orientation, data migration, and a phased changeover strategy.')
doc.add_paragraph('Installation components include:')
doc.add_paragraph('• Environment preparation (server, database, network)', style='List Bullet')
doc.add_paragraph('• Application deployment and configuration', style='List Bullet')
doc.add_paragraph('• User training session for system operations', style='List Bullet')
doc.add_paragraph('• Data migration from legacy sources', style='List Bullet')

doc.add_paragraph('Changeover strategies:')
doc.add_paragraph('• Parallel changeover: run the new system alongside the old system until verified.', style='List Bullet')
doc.add_paragraph('• Phased changeover: introduce modules gradually to reduce risk.', style='List Bullet')
doc.add_paragraph('• Recommended strategy: Parallel changeover to allow validation of credit scores before full cutover.', style='List Bullet')

doc.add_heading('8.5 Maintenance', level=2)
doc.add_paragraph(
    'Maintenance recommendations include regular system updates, security patching, data backups, and user support. A maintenance plan should define routine tasks, monitoring processes, and a support escalation path.')
doc.add_paragraph('Recommended maintenance strategies:')
doc.add_paragraph('• Preventive maintenance: regular updates and vulnerability scanning.', style='List Bullet')
doc.add_paragraph('• Corrective maintenance: fix bugs and respond to issues quickly.', style='List Bullet')
doc.add_paragraph('• Adaptive maintenance: update the system for new credit scoring rules or regulations.', style='List Bullet')

doc.add_heading('8.6 Recommendations for future/further development', level=2)
doc.add_paragraph(
    'Future development may include machine learning-driven credit risk models, mobile application support, automated data import from financial institutions, and expanded analytics dashboards. ' 
    'These enhancements can improve predictive accuracy and user experience.')
doc.add_heading('8.7 Conclusion', level=2)
doc.add_paragraph(
    'This implementation and testing chapter confirms that the Credit Score System is ready for deployment and operation. ' 
    'The design supports secure data processing, controlled installation, and ongoing maintenance while allowing future growth and improvements.')

doc.save('Credit_Score_System_Dissertation.docx')
print('Created Credit_Score_System_Dissertation.docx')