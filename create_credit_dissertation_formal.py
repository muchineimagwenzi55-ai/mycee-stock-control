from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import os

OUTPUT_DOC = 'Credit_Score_System_Dissertation_Formal.docx'
DIAGRAM_DIR = 'dissertation_diagrams'

if not os.path.exists(DIAGRAM_DIR):
    os.makedirs(DIAGRAM_DIR)

font = ImageFont.load_default()


def draw_diagram(path, title, elements, arrows=None):
    width, height = 1200, 700
    background = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(background)

    draw.rectangle([20, 20, width - 20, height - 20], outline='black', width=2)
    draw.text((40, 30), title, fill='black', font=font)

    for element in elements:
        x, y, w, h, label = element
        draw.rectangle([x, y, x + w, y + h], outline='black', width=2)
        draw.text((x + 10, y + 10), label, fill='black', font=font)

    if arrows:
        for arrow in arrows:
            x1, y1, x2, y2, label = arrow
            draw.line([x1, y1, x2, y2], fill='black', width=2)
            draw.polygon([ (x2, y2), (x2-10, y2-10), (x2-10, y2+10) ], fill='black')
            if label:
                mx = (x1 + x2) / 2
                my = (y1 + y2) / 2
                draw.text((mx + 5, my - 10), label, fill='black', font=font)

    background.save(path)


def create_diagrams():
    draw_diagram(
        os.path.join(DIAGRAM_DIR, 'context_diagram.png'),
        'Figure 7.1: Context Diagram',
        [
            (850, 100, 250, 80, 'Borrower'),
            (850, 250, 250, 80, 'Credit Analyst'),
            (850, 400, 250, 80, 'Data Provider'),
            (300, 250, 300, 120, 'Credit Score System')
        ],
        [
            (600, 310, 850, 140, 'Request\nInformation'),
            (850, 190, 600, 310, 'Borrower Data'),
            (600, 360, 850, 280, 'Score Request'),
            (850, 330, 600, 360, 'Score Report'),
            (850, 420, 600, 360, 'External Data')
        ]
    )

    draw_diagram(
        os.path.join(DIAGRAM_DIR, 'dfd.png'),
        'Figure 7.2: Data Flow Diagram',
        [
            (50, 250, 220, 120, 'Borrower Data Source'),
            (950, 250, 220, 120, 'Credit Bureau'),
            (390, 60, 420, 140, 'Validation Process'),
            (390, 260, 420, 140, 'Scoring Process'),
            (390, 460, 420, 140, 'Reporting Process'),
            (640, 250, 220, 120, 'Database')
        ],
        [
            (270, 310, 390, 120, 'Financial Data'),
            (820, 310, 810, 320, 'Credit Data'),
            (570, 160, 570, 260, 'Validated Data'),
            (570, 340, 570, 460, 'Score Inputs'),
            (570, 540, 570, 460, 'Report Params'),
            (760, 310, 760, 310, 'Store/Fetch')
        ]
    )

    draw_diagram(
        os.path.join(DIAGRAM_DIR, 'eer_diagram.png'),
        'Figure 7.3: EER Diagram',
        [
            (50, 120, 240, 100, 'Borrower'),
            (380, 120, 240, 100, 'FinancialRecord'),
            (710, 120, 240, 100, 'CreditScore'),
            (380, 320, 240, 100, 'ScoreRule'),
            (710, 320, 240, 100, 'AuditLog')
        ],
        [
            (290, 170, 380, 170, '1-to-many'),
            (620, 170, 710, 170, '1-to-many'),
            (500, 220, 500, 320, None),
            (820, 220, 820, 320, 'Logs')
        ]
    )

    draw_diagram(
        os.path.join(DIAGRAM_DIR, 'sequence_diagram.png'),
        'Figure 7.4: Sequence Diagram',
        [
            (100, 150, 200, 80, 'Borrower'),
            (400, 150, 200, 80, 'Web UI'),
            (700, 150, 200, 80, 'Scoring Service'),
            (1000, 150, 200, 80, 'Database')
        ],
        [
            (200, 190, 400, 190, 'Submit Application'),
            (500, 190, 700, 190, 'Validate & Score'),
            (900, 190, 1000, 190, 'Fetch History'),
            (1000, 250, 700, 250, 'Return Data'),
            (700, 250, 400, 250, 'Score Result'),
            (400, 250, 200, 250, 'Display Score')
        ]
    )

    draw_diagram(
        os.path.join(DIAGRAM_DIR, 'package_diagram.png'),
        'Figure 7.5: Package Diagram',
        [
            (60, 100, 220, 100, 'auth'),
            (350, 100, 220, 100, 'borrower'),
            (640, 100, 220, 100, 'scoring'),
            (60, 300, 220, 100, 'reporting'),
            (350, 300, 220, 100, 'persistence')
        ],
        [
            (190, 200, 190, 300, 'secure data'),
            (470, 200, 470, 300, 'score results'),
            (360, 350, 470, 350, 'save/load')
        ]
    )


create_diagrams()

doc = Document()

# Title and chapter headings

doc.add_heading('Credit Score System Dissertation', level=1)

doc.add_heading('Chapter 7: Design', level=1)

doc.add_heading('7.1 Introduction', level=2)
doc.add_paragraph(
    'Chapter 7 presents the design of the proposed Credit Score System. ' 
    'It provides a structured description of how the system will function, the ' 
    'software and hardware architecture, database organization, program structure, ' 
    'interface layout, and security measures. This chapter lays the foundation for ' 
    'the subsequent implementation and testing phases.')

# System Design

doc.add_heading('7.2 System Design', level=2)
doc.add_paragraph(
    'The Credit Score System is intended to evaluate borrower creditworthiness by ' 
    'collecting relevant financial information, applying defined scoring rules, and ' 
    'producing a standardized credit score. It is designed to support secure data ' 
    'entry, reliable score generation, and comprehensive reporting for credit analysts.')

doc.add_paragraph('The major system functions are:')
doc.add_paragraph('• Borrower registration and profile management', style='List Bullet')
doc.add_paragraph('• Financial data collection and validation', style='List Bullet')
doc.add_paragraph('• Credit scoring algorithm execution', style='List Bullet')
doc.add_paragraph('• Report generation and decision support', style='List Bullet')
doc.add_paragraph('• Audit logging and security monitoring', style='List Bullet')

doc.add_heading('7.2.1 Context Diagram and DFD of the Proposed System', level=3)
doc.add_paragraph(
    'The context diagram defines system boundaries and external actors. The system ' 
    'receives data from borrowers and third-party sources, processes scoring requests, ' 
    'and delivers credit reports and scores to authorized users.')
doc.add_paragraph(
    'The data flow diagram elaborates the internal processes, illustrating how data ' 
    'is validated, how the credit score is computed, and how reports are produced. ' 
    'The DFD supports understanding of data movement and information processing.')

doc.add_paragraph('The DFD includes the following levels:')
doc.add_paragraph('• Level 0: System context with borrower and external data provider inputs.', style='List Bullet')
doc.add_paragraph('• Level 1: Primary processes for validation, scoring, and reporting.', style='List Bullet')
doc.add_paragraph('• Level 2: Detailed subprocesses for risk calculation, data storage, and response generation.', style='List Bullet')

doc.add_paragraph('The figures below illustrate the system context and major data flows.')

doc.add_picture(os.path.join(DIAGRAM_DIR, 'context_diagram.png'), width=Inches(6))
doc.add_paragraph('Figure 7.1: Context Diagram of the proposed Credit Score System.', style='Caption')
doc.add_picture(os.path.join(DIAGRAM_DIR, 'dfd.png'), width=Inches(6))
doc.add_paragraph('Figure 7.2: Data Flow Diagram for the proposed Credit Score System.', style='Caption')

# Architectural design

doc.add_heading('7.3 Architectural Design', level=2)
doc.add_paragraph(
    'The system architecture employs a three-tier design comprising a presentation ' 
    'layer, an application layer, and a data layer. The presentation layer offers a ' 
    'web-based user interface for borrowers and credit analysts. The application layer ' 
    'implements business logic for scoring, validation, and reporting. The data layer ' 
    'maintains persistent storage and supports secure retrieval of historical data.')
doc.add_paragraph('Core architectural components are:')
doc.add_paragraph('• Presentation layer: browser-based UI and dashboards.', style='List Bullet')
doc.add_paragraph('• Business logic layer: scoring engine and validation services.', style='List Bullet')
doc.add_paragraph('• Data layer: relational database and data access services.', style='List Bullet')
doc.add_paragraph('• Integration layer: secure interfaces to external credit bureaus and data sources.', style='List Bullet')

doc.add_paragraph('This architecture supports modular development, scalability, and separation of concerns, allowing the scoring engine to be updated independently of the presentation layer.')

doc.add_heading('7.4 Physical Design', level=2)
doc.add_paragraph(
    'The physical design describes the runtime deployment of the system on hardware and network infrastructure. ' 
    'The proposed deployment uses a cloud-hosted application server, an external relational database service, ' 
    'and user devices that access the system through secured web browsers.')
doc.add_paragraph('Software and hardware interactions are as follows:')
doc.add_paragraph('• Client devices communicate with the web server over HTTPS.', style='List Bullet')
doc.add_paragraph('• The application server processes business logic and interacts with the database.', style='List Bullet')
doc.add_paragraph('• The database server stores borrower and scoring records with secure access controls.', style='List Bullet')
doc.add_paragraph('• Third-party services provide additional credit and financial data through encrypted API connections.', style='List Bullet')

doc.add_heading('7.5 Database Design', level=2)
doc.add_paragraph(
    'The relational database schema is designed to support borrower information, financial records, credit scores, scoring rules, and audit logs. ' 
    'The schema emphasizes normalization, referential integrity, and efficient querying.')
doc.add_paragraph('Primary database tables include:')
doc.add_paragraph('• Borrower', style='List Bullet')
doc.add_paragraph('• FinancialRecord', style='List Bullet')
doc.add_paragraph('• CreditScore', style='List Bullet')
doc.add_paragraph('• ScoreRule', style='List Bullet')
doc.add_paragraph('• AuditLog', style='List Bullet')
doc.add_paragraph('The enhanced entity-relationship diagram below documents principal entities and their relationships.')
doc.add_picture(os.path.join(DIAGRAM_DIR, 'eer_diagram.png'), width=Inches(6))
doc.add_paragraph('Figure 7.3: Enhanced Entity-Relationship Diagram for the Credit Score System.', style='Caption')

doc.add_heading('7.6 Program Design', level=2)
doc.add_paragraph(
    'Program design defines the internal software structure, package organization, and interactions among classes and services. This design promotes maintainability, reuse, and clear separation of responsibilities.')
doc.add_paragraph('Suggested package structure includes:')
doc.add_paragraph('• auth – handles authentication and authorization.', style='List Bullet')
doc.add_paragraph('• borrower – manages borrower profiles and data.', style='List Bullet')
doc.add_paragraph('• scoring – implements credit score calculation logic.', style='List Bullet')
doc.add_paragraph('• reporting – generates score and risk reports.', style='List Bullet')
doc.add_paragraph('• persistence – manages database access and data storage.', style='List Bullet')
doc.add_paragraph('The package diagram below illustrates the relationships among primary software modules.')
doc.add_picture(os.path.join(DIAGRAM_DIR, 'package_diagram.png'), width=Inches(6))
doc.add_paragraph('Figure 7.4: Package Diagram showing major application modules.', style='Caption')
doc.add_paragraph('A sequence diagram for the credit scoring process is shown in Figure 7.5.')
doc.add_picture(os.path.join(DIAGRAM_DIR, 'sequence_diagram.png'), width=Inches(6))
doc.add_paragraph('Figure 7.5: Sequence Diagram for borrower credit scoring interaction.', style='Caption')

doc.add_heading('7.7 Interface Design', level=2)
doc.add_paragraph(
    'Interface design defines the user interaction and ensures that the system is intuitive and accessible. ' 
    'Menus, data entry forms, and reports are arranged to support efficient workflow and accurate data capture.')
doc.add_heading('7.7.1 Menu Design', level=3)
doc.add_paragraph('The main menu structure is organized into logical categories for users.')
doc.add_heading('7.7.1.1 Main Menu', level=4)
doc.add_paragraph('The main menu provides access to:')
doc.add_paragraph('• Dashboard', style='List Bullet')
doc.add_paragraph('• Borrower Management', style='List Bullet')
doc.add_paragraph('• Credit Scoring', style='List Bullet')
doc.add_paragraph('• Reports', style='List Bullet')
doc.add_paragraph('• Administration', style='List Bullet')
doc.add_heading('7.7.1.2 Sub-menus', level=4)
doc.add_paragraph('Sub-menu options include:')
doc.add_paragraph('• Add New Borrower', style='List Bullet')
doc.add_paragraph('• Upload Financial Data', style='List Bullet')
doc.add_paragraph('• View Score History', style='List Bullet')
doc.add_paragraph('• Generate Report', style='List Bullet')
doc.add_heading('7.7.2 Input Design', level=3)
doc.add_paragraph(
    'Input design ensures that each form collects the required data with validation and clarity. ' 
    'Each input field is labelled clearly and includes data validation rules to reduce entry errors.')
doc.add_paragraph('Input forms include:')
doc.add_paragraph('• Borrower registration form', style='List Bullet')
doc.add_paragraph('• Income and employment details form', style='List Bullet')
doc.add_paragraph('• Loan and liability history form', style='List Bullet')
doc.add_paragraph('• Credit bureau data import form', style='List Bullet')
doc.add_heading('7.7.3 Output Design', level=3)
doc.add_paragraph(
    'Output design specifies the reports and summaries produced by the system. Outputs are designed for clarity, decision support, and auditability.')
doc.add_paragraph('Output reports include:')
doc.add_paragraph('• Credit score report', style='List Bullet')
doc.add_paragraph('• Borrower risk summary', style='List Bullet')
doc.add_paragraph('• Credit history analysis', style='List Bullet')
doc.add_paragraph('• Audit log report', style='List Bullet')
doc.add_heading('7.8 Pseudo Code', level=2)
doc.add_paragraph(
    'The following pseudo code describes the core scoring algorithm and data validation flow:')
doc.add_paragraph(
    '1. Receive borrower profile and financial input.\n'
    '2. Validate required fields and data consistency.\n'
    '3. Retrieve historical credit and repayment records.\n'
    '4. Apply scoring rules to compute risk points.\n'
    '5. Aggregate risk points into a final credit score.\n'
    '6. Store the score and generate a report.\n'
    '7. Present the result to the authorized user.\n'
)

doc.add_heading('7.9 Security Design', level=2)
doc.add_paragraph(
    'Security design is vital for protecting borrower data and ensuring regulatory compliance. ' 
    'The system employs multiple layers of defense including physical, network, and operational security.')
doc.add_heading('7.9.1 Physical Security', level=3)
doc.add_paragraph(
    'Physical security controls protect the data center and operational workstations. ' 
    'Measures include restricted access to servers, locked facilities, CCTV monitoring, and secure disposal of sensitive documents.')
doc.add_heading('7.9.2 Network Security', level=3)
doc.add_paragraph(
    'Network security ensures all communications are encrypted and unauthorized access is blocked. ' 
    'The system uses HTTPS for browser access, VPNs for administrative access, firewalls, and intrusion detection systems.')
doc.add_heading('7.9.3 Operational Security', level=3)
doc.add_paragraph(
    'Operational security includes role-based access control, strong authentication, audit logging, ' 
    'regular security reviews, and employee awareness training. Only authorized staff may access sensitive functionality.')
doc.add_heading('7.10 Conclusion', level=2)
doc.add_paragraph(
    'Chapter 7 has presented a comprehensive design for the proposed Credit Score System. ' 
    'The design supports a robust, secure, and scalable implementation that meets the needs of borrowers, analysts, and regulatory stakeholders.')

doc.add_heading('Chapter 8: Implementation and Testing', level=1)
doc.add_heading('8.1 Introduction', level=2)
doc.add_paragraph(
    'Chapter 8 outlines the implementation process, testing regime, installation strategy, maintenance planning, and future developments for the Credit Score System.')
doc.add_heading('8.2 Coding', level=2)
doc.add_paragraph(
    'The implementation employs modular code organization with clear separation between user interface, business logic, and persistence. ' 
    'Key techniques include input validation, error handling, secure authentication, and reusable services.')
doc.add_heading('8.3 Testing', level=2)
doc.add_paragraph(
    'Testing covers functional requirements, user acceptance, and security assurance. ' 
    'The system must demonstrate reliable borrower registration, accurate score calculation, and enforcement of access control policies.')
doc.add_paragraph('Representative test cases:')
doc.add_paragraph('• Create borrower profile with valid data', style='List Bullet')
doc.add_paragraph('• Reject invalid or incomplete borrower entries', style='List Bullet')
doc.add_paragraph('• Verify credit score generation for different borrower profiles', style='List Bullet')
doc.add_paragraph('• Confirm unauthorized users cannot access restricted features', style='List Bullet')
doc.add_paragraph(
    'Security testing should include authentication enforcement, input validation, and protection against unauthorized data access. ' 
    'Insert screenshots of test results and security checks to demonstrate compliance with the security design described in Chapter 7.')
doc.add_heading('8.4 Installation', level=2)
doc.add_paragraph(
    'Installation includes software deployment, user training, data migration, and changeover planning. ' 
    'A structured rollout ensures that users adopt the system smoothly and that legacy data are preserved.')
doc.add_paragraph('Key installation activities include:')
doc.add_paragraph('• Preparing the deployment environment and configuring servers', style='List Bullet')
doc.add_paragraph('• Installing application software and database components', style='List Bullet')
doc.add_paragraph('• Conducting user training sessions and documentation delivery', style='List Bullet')
doc.add_paragraph('• Migrating legacy borrower data and validating consistency', style='List Bullet')
doc.add_paragraph('• Executing the chosen changeover strategy', style='List Bullet')
doc.add_paragraph('Recommended changeover strategy:')
doc.add_paragraph('• Parallel changeover: operate the new system alongside the existing system until verification is complete.', style='List Bullet')
doc.add_paragraph('This approach minimizes risk and allows comparison of outputs before full transition.', style='List Bullet')
doc.add_heading('8.5 Maintenance', level=2)
doc.add_paragraph(
    'Maintenance planning ensures the system remains reliable and secure over time. ' 
    'The maintenance strategy should include scheduled updates, backups, monitoring, and issue resolution procedures.')
doc.add_paragraph('Recommended maintenance strategies:')
doc.add_paragraph('• Preventive maintenance: apply patches, updates, and security fixes regularly.', style='List Bullet')
doc.add_paragraph('• Corrective maintenance: address identified defects and operational issues promptly.', style='List Bullet')
doc.add_paragraph('• Adaptive maintenance: update the system to meet new scoring rules, regulations, or business requirements.', style='List Bullet')
doc.add_heading('8.6 Recommendations for Future Development', level=2)
doc.add_paragraph(
    'Future enhancements may include machine learning-driven risk models, expanded mobile access, automated data feeds, and advanced analytics dashboards. ' 
    'These developments would improve predictive accuracy, user convenience, and decision support.')
doc.add_heading('8.7 Conclusion', level=2)
doc.add_paragraph(
    'Chapter 8 confirms that the proposed Credit Score System can be implemented with a disciplined coding and testing approach, installed with minimal disruption, and maintained effectively. ' 
    'It also identifies opportunities for future improvements to support evolving business needs.')

doc.save(OUTPUT_DOC)
print(f'Created {OUTPUT_DOC}')
